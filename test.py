import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import unicodedata
import re
import os

# Spații invizibile (zero-width, non-breaking etc.)
INVISIBLE_SPACES = set(['\u00A0', '\u200B', '\u2009', '\u202F'])

# Caractere de întrâmpuire automat cu minus
DASH_REPLACEMENTS = {'—': '-', '–': '-'}

# Funcție de detectare spații invizibile
def is_invisible_space(char):
    return char in INVISIBLE_SPACES

class DocxCheckerApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Detector spații invizibile și liniuțe speciale .docx")
        self.text_display = tk.Text(master, wrap=tk.WORD, height=10, width=80)
        self.text_display.pack(padx=10, pady=10)
        self.load_button = tk.Button(master, text="Încarcă fișier .docx", command=self.load_docx)
        self.load_button.pack(pady=5)
        self.file_path = None
        self.doc = None

    def load_docx(self):
        self.file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not self.file_path:
            return
        self.doc = Document(self.file_path)
        self.replace_dashes()
        self.clean_trailing_invisible_spaces()

    def replace_dashes(self):
        modified = False
        for para in self.doc.paragraphs:
            old_text = para.text
            new_text = old_text
            for dash_char, replacement in DASH_REPLACEMENTS.items():
                if dash_char in new_text:
                    new_text = new_text.replace(dash_char, replacement)
                    modified = True
                    print(f"Înlocuit '{dash_char}' cu '-' în: {old_text.strip()}")
            para.text = new_text
        if modified:
            self.doc.save(self.file_path)
            print("Toate caracterele — și – au fost înlocuite cu - și fișierul a fost salvat.")

    def clean_trailing_invisible_spaces(self):
        modified = False
        for para in self.doc.paragraphs:
            old_text = para.text
            trimmed = old_text.rstrip()
            removed = False
            while trimmed and is_invisible_space(trimmed[-1]):
                print(f"Eliminat caracter invizibil de la sfârșit: {repr(trimmed[-1])} (U+{ord(trimmed[-1]):04X}) în: {old_text.strip()}")
                trimmed = trimmed[:-1]
                removed = True
            if removed:
                para.text = trimmed
                modified = True
        if modified:
            self.doc.save(self.file_path)
            print("Spațiile invizibile de la sfârșitul rândurilor au fost eliminate și fișierul a fost salvat.")
            self.show_cleaned_message()
        else:
            messagebox.showinfo("Rezultat", "Nu au fost găsite spații invizibile la final sau liniuțe speciale.")

    def show_cleaned_message(self):
        self.text_display.delete('1.0', tk.END)
        self.text_display.insert(tk.END, "Documentul a fost curățat cu succes. Verifică consola pentru detalii.")

if __name__ == '__main__':
    root = tk.Tk()
    app = DocxCheckerApp(root)
    root.mainloop()
